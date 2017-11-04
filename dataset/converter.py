
import os, sys, shutil
import docx

def docx_to_str(filename):
##    global path_to
    doc = docx.Document(filename)
    fullText = []
##    if (len(doc.tables) == 0):
##        print('tables 0', filename)
##        shutil.copy2(filename, path_to)
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def docx_to_txt(filename, path_to = None):
    txt_name = os.path.splitext(os.path.basename(filename))[0] + '.txt'
    if path_to is not None:
        txt_name = os.path.join(path_to, txt_name)
    s = docx_to_str(filename)
    with open(txt_name, 'w') as f:
        f.write(s)

path_from = r'D:\Temp\docx\docx0tables'
path_to = r'D:\Temp\docx\txt'
if not os.path.isdir(path_to):
    os.makedirs(path_to)
    
for f in os.listdir(path_from):
    if f.lower().endswith('.docx'):
        print(f)
        try:
            docx_to_txt(os.path.join(path_from, f), path_to)
        except:
            print('remove', f)
            t = os.path.splitext(f)[0] + '.txt'
            os.remove(os.path.join(path_from, f))
            os.remove(os.path.join(path_to, t))


            
